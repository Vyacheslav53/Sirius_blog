
# Устанавливаем библиотеки
# В терминале пишем:
# pip install requests
# pip install selenium
# pip install pandas

# Импортируем библиотеки
import pandas as pd
import requests
import time
from selenium.webdriver.chrome.service import Service
from docx import Document

from selenium import webdriver
from selenium.webdriver.common.by import By

from bs4 import BeautifulSoup

from PIL import Image
from io import BytesIO

from dotenv import load_dotenv
import openai
import os

from openai import OpenAI

load_dotenv()
# API-key
openai.api_key = os.environ.get("OPENAI_API_KEY")

client = OpenAI()

# Создаем путь к драйверу chrome
service = Service('C:/Users/chromedriver.exe')

# Создание экземпляра веб-драйвера Chrome
driver = webdriver.Chrome(service = service)

# Ключевые слова для поиска новостей
keywords = 'важность раннего обучения детей арифметике скорочтению'

# Формирование URL для поиска
start_url = "https://dzen.ru/news/search"
# Открытие страницы
driver.get(start_url)

search_input = driver.find_element(By.XPATH, '//input[@placeholder="Найти в Новостях"]')

# Указываем поиск по ключевым словам
search_input.send_keys(keywords)
time.sleep(10)

# Находим и нажимаем кнопку "Найти"
search_button = driver.find_element(By.XPATH, '//button[@aria-label="Кнопка Найти"]')
search_button.click()
time.sleep(10)

# Поиск всех новостных блоков
news_blocks = driver.find_elements(By.CSS_SELECTOR, '.news-search-content__block.news-search-content__block_doc')

# Создание списка для хранения данных
news_data_link = []

# Обработка каждого новостного блока для получения ссылки на новость
for i, news in enumerate(news_blocks, start=1):
    
    # Ссылка
    link_tag = news.find_element(By.TAG_NAME, 'a')
    link = link_tag.get_attribute('href') if link_tag else 'Без ссылки'

    # Добавление данных в список
    news_data_link.append(link)

# Закрытие веб-драйвера
driver.quit()

# Запись ссылок в csv файл
# Создание датафрейма
df = pd.DataFrame(news_data_link, columns=['Ссылка'])

# Запись данных в csv файл
df.to_csv('news_data.csv', index=False, encoding='utf-8-sig')

# Чтение ссылок из файла
df = pd.read_csv('news_data.csv')
news_data_link = df['Ссылка'].tolist()

# Парсинг текстов новостей по ссылкам
news_text = ""

# Заголовки для имитации запроса от браузера
headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
}

# Функция парсинга новостей
def get_text_from_url(news_url):
    # Выполнение запроса
    response = requests.get(news_url, headers=headers)

    # Проверка успешного выполнения запроса
    if response.status_code == 200:
        # Парсинг HTML-контента
        soup = BeautifulSoup(response.content, 'html.parser')

        # Поиск основного текста
        main_text_div = soup.find('div', class_='entry-content')

        if main_text_div:
            # Извлечение текстового содержимого
            return main_text_div.get_text(separator='\n', strip=True)            
        else:
            return ""
    else:
        return ""

# Проходим по всем ссылкам и собираем текст
for news_url in news_data_link:
    page_text = get_text_from_url(news_url)
    news_text += page_text + "\n\n"  # Добавляем текст из каждой страницы

# Сохраняем текст новостей в файл docx

# Создание нового документа
doc = Document()
doc.add_paragraph(news_text)

# Сохранение документа в файл
doc.save('news_text.docx')


# Чтение текста новостей из файла docx
doc_news = Document('news_text.docx')
# Переменная для хранения текста
news_text = ""
# Проходим по всем параграфам в документе и добавляем их в переменную
for paragraph in doc_news.paragraphs:
    news_text += paragraph.text + "\n"  # Добавляем текст параграфа с новой строки


# Генерация текста для блога из текста новостей

# Текст для system_blog для подачи в GPT
doc = Document('PROMPT_bloger.docx')
system_blog = ""
for paragraph in doc.paragraphs:
    system_blog += paragraph.text + "\n"

completion = client.chat.completions.create(
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": system_blog},
        {
            "role": "user",
            "content": f'''Проанализируй информацию из следующих новостей и составь краткий интересный текст для блога. Новости:
            {news_text}'''
        }
    ]
)

blog_text = (completion.choices[0].message).content

# Сохранение текста для блога в файл docs
doc_blog = Document()
doc_blog.add_paragraph(blog_text)

# Сохранение документа в файл
doc_blog.save('blog_text.docx')

# Чтение текста для блога из файла
doc = Document('blog_text.docx')
blog_text = ""
for paragraph in doc.paragraphs:
    blog_text += paragraph.text + "\n"


# Генерация текста для создания изображения
# Текст для system_copywriter для подачи в GPT
doc = Document('PROMPT_copywriter.docx')
system_copywriter = ""
for paragraph in doc.paragraphs:
    system_copywriter += paragraph.text + "\n"


# I вариант с gpt-3.5-turbo
completion = client.chat.completions.create(
    model="gpt-3.5-turbo",
    messages=[
        {"role": "system", "content": system_copywriter},
        {
            "role": "user",
            "content": f'''Проанализируй информацию из текста блога и составь краткий яркий запоминающийся лозунг для создания изображения. Текст блога:
            {blog_text}'''
        }
    ]
)
slogan = (completion.choices[0].message).content

doc = Document()
doc.add_paragraph(slogan)
doc.save('slogan.docx')

# Чтение слогана из файла 
doc = Document('slogan.docx')
slogan = ""
for paragraph in doc.paragraphs:
    slogan += paragraph.text + "\n"

# II вариант с gpt-4o-mini
completion = client.chat.completions.create(
    model="gpt-4o-mini",
    messages=[
        {"role": "system", "content": system_copywriter},
        {
            "role": "user",
            "content": f'''Проанализируй информацию из текста блога и составь краткий яркий запоминающийся лозунг для создания изображения. Текст блога:
            {blog_text}'''
        }
    ]
)
slogan1 = (completion.choices[0].message).content

doc = Document()
doc.add_paragraph(slogan1)
doc.save('slogan1.docx')

# Чтение слогана из файла 
doc = Document('slogan1.docx')
slogan1 = ""
for paragraph in doc.paragraphs:
    slogan1 += paragraph.text + "\n"

# Генерация картинки из текста слогана
# 2 варианта промптов

response = client.images.generate(
  model="dall-e-3",
  prompt=slogan,
  size="1024x1024",
  quality="standard",
  n=1,
)
image_url = response.data[0].url

response = client.images.generate(
  model="dall-e-3",
  prompt=f'Инфографика для текста: {slogan1}',
  size="1024x1024",
  quality="standard",
  n=1,
)
image_url_1 = response.data[0].url

# Функция сохранения изображения в файл png
def download_and_save_image(url, save_path):
    # Загрузка изображения по URL
    response = requests.get(url)
    response.raise_for_status()  # Проверка успешности запроса
    # Открытие изображения с помощью PIL
    image = Image.open(BytesIO(response.content))
    # Сохранение изображения в формате PNG
    image.save(save_path, format='PNG')

# Пути для сохранения картинки
save_path = 'image.png'
save_path_1 = 'image_1.png'

# Сохранение картинки в файл png
download_and_save_image(image_url, save_path)
download_and_save_image(image_url_1, save_path_1)